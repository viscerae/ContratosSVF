import FreeSimpleGUI as sg
import pandas as pd
from docx import Document
from datetime import datetime
import os  
from docx.shared import Pt
from num2words import num2words
import locale

sg.theme("dark")
locale.setlocale(locale.LC_TIME, "pt_PT")

layout = [
    [sg.Text("Excel"), 
     sg.FileBrowse("Escolher Ficheiro Excel", target="input_excel", button_color=('white', 'blue')), 
     sg.Input(key="input_excel")],
    [sg.Text("Word"), 
     sg.FileBrowse("Escolher Ficheiro Word", target="input_word",  button_color=('white', 'blue')), 
     sg.Input(key="input_word")],
    [sg.Button("Gerar Contratos Completos", button_color=('white', 'blue'))]
]

window = sg.Window("Gerador de Contrato", layout)

def preencher_contrato(excel_path, word_path, folder_name, today):
    print(excel_path)
    try:
        excel = pd.read_excel(excel_path)
        
        for _, row in excel.iterrows():
            print(row)
    ######################################## DATA EXTRACTION ############################################
            word_doc = Document(word_path)  
            nome = row.iloc[1]
            estado_civil = row.iloc[2]
            morada = f"{row.iloc[3]}, {row.iloc[4]}" 
            naturalidade = row.iloc[5]
            tipo_id = str(row.iloc[6])
            nr_id = str(row.iloc[7])
            validade_raw = str(row.iloc[8])
            validade_raw = validade_raw[:10]
            val = pd.to_datetime(validade_raw.strip(), errors='coerce').strftime('%d/%m/%Y')  
            nif = row.iloc[9]
            niss = row.iloc[10]
            categoria = str(row.iloc[11])
            func = str(row.iloc[12])
            horas_semanal = str(row.iloc[13])
            horas_diario = str(row.iloc[14])
            renum = row.iloc[15]
            ext = num2words(renum, lang='pt')
            inic_contrato = str(row.iloc[16])
            inic_contrato = inic_contrato[:10]
            datacontrato = pd.to_datetime(inic_contrato.strip(), errors='coerce').strftime('%d/%m/%Y')  
    #####################################################################################################
       
    ########################################## TEXT REPLACEMENT #########################################        
            replacement_text = (
                f"{nome}, {estado_civil}, natural de {naturalidade}, residente na {morada}., portador(a) do {tipo_id} n.º {nr_id}, válido até {val}, contribuinte fiscal n.º {nif},e NISS {niss} de ora em diante designada apenas por ""Trabalhador"";"
            )
            replacement_text2 = (
                f"categoria de {categoria}, para que desempenhe, sob as ordens e direcção daquela, as funções inerentes àquela categoria e, designadamente: {func}."
            )
            replacement_text3 = (
                f"O horário de trabalho em vigor na Empresa é de {horas_semanal} horas semanais, com {horas_diario} horas diárias a prestar de segunda a sexta-feira entre as 09:00 horas e as 18:00  horas, com intervalo de uma hora para almoço."
            )
            replacement_text4 = (
                f"Como contrapartida pela prestação de trabalho prevista neste contrato a Empresa pagará ao Trabalhador, mediante transferência bancária ou, excecionalmente e por motivos de necessidade operacional, através de cheque bancário uma remuneração mensal ilíquida de {renum:.2f} € ({ext} euros)."
            )
            replacement_text5 = (
                f"Feito em duas vias, em Lisboa, no dia {datacontrato}"
            )
            replacement_text6 = (
                f"O presente contrato é celebrado sem termo, produzindo efeitos a partir do dia {datacontrato}"
            )
            
            for paragraph in word_doc.paragraphs:
                if '[TRABALHADOR]'in paragraph.text:
                    paragraph.clear()
                    run = paragraph.add_run(replacement_text)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
            for paragraph in word_doc.paragraphs:
                if "[CATEGORIA]" in paragraph.text:
                    paragraph.clear()
                    run = paragraph.add_run(replacement_text2)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
            for paragraph in word_doc.paragraphs:
                if "[HORAS]" in paragraph.text:
                    paragraph.clear()
                    run = paragraph.add_run(replacement_text3)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
            for paragraph in word_doc.paragraphs:
                if "[RENUM]" in paragraph.text:
                    paragraph.clear()
                    run = paragraph.add_run(replacement_text4)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
            for paragraph in word_doc.paragraphs:
                if "[DATAA]" in paragraph.text:
                    paragraph.clear()
                    run = paragraph.add_run(replacement_text5)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
            for paragraph in word_doc.paragraphs:
                if "[INITCONT]" in paragraph.text:
                    paragraph.clear()
                    run = paragraph.add_run(replacement_text6)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
            
    #####################################################################################################           
 
            output_path = os.path.join(folder_name, f"{nome}.docx")
            word_doc.save(output_path)
            print(f"Documento gerado: {output_path}")
            
        sg.popup("Documentos gerados com sucesso!", "Arquivos foram salvos em uma nova pasta.")

    except Exception as e:
        print(f"An error occurred: {e}")
        sg.popup(f"Error: {e}")
        
    ######################################### WINDOW LOOP ###############################################
    
while True:
    event, values = window.read()
    if event == sg.WIN_CLOSED:     
        break                    
    elif event == "Gerar Contratos Completos":      
        today = datetime.today().strftime("%d de %B de %Y")
        desktop_path = os.path.expanduser("~/Desktop")
        folder_name = os.path.join(desktop_path, f"Contratos - {today}")

        os.makedirs(folder_name, exist_ok=True)

        
        excel_path = values["input_excel"]
        word_path = values["input_word"]
        
        preencher_contrato(excel_path, word_path, folder_name, today)

window.close()
    #####################################################################################################
